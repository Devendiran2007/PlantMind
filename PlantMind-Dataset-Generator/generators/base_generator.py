import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.pdfgen import canvas

# --- PDF Page Numbering Canvas ---
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_decorations(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica-Bold", 8)
        self.setFillColor(colors.HexColor("#1A365D")) # Deep Navy
        
        # Header
        self.drawString(54, 795, "STEELFORGE INDUSTRIES PVT LTD")
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#4A5568")) # Charcoal
        self.drawRightString(541, 795, "PlantMind Industrial Document Management System")
        
        # Header Line
        self.setStrokeColor(colors.HexColor("#CBD5E0")) # Light grey
        self.setLineWidth(0.75)
        self.line(54, 787, 541, 787)
        
        # Footer Line
        self.line(54, 50, 541, 50)
        
        # Footer
        self.drawString(54, 38, "CONFIDENTIAL - TECHNICAL INTEGRITY DATA")
        self.drawRightString(541, 38, f"Page {self._pageNumber} of {page_count}")
        self.restoreState()

# --- DOCX Helper functions for borders and shading ---
def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

class BaseGenerator:
    def __init__(self, output_dir):
        self.output_dir = Path(output_dir)
        self.pdf_dir = self.output_dir / "PDF"
        self.docx_dir = self.output_dir / "DOCX"
        self.pdf_dir.mkdir(parents=True, exist_ok=True)
        self.docx_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_pdf(self, doc_data, category_name):
        category_pdf_dir = self.pdf_dir / category_name
        category_pdf_dir.mkdir(parents=True, exist_ok=True)
        pdf_path = category_pdf_dir / f"{doc_data['doc_id']}.pdf"
        
        # Setup document layout (A4 with 0.75in margins)
        doc = SimpleDocTemplate(
            str(pdf_path),
            pagesize=A4,
            leftMargin=54,
            rightMargin=54,
            topMargin=72,
            bottomMargin=72
        )
        
        styles = getSampleStyleSheet()
        
        # Modify existing styles to avoid conflicts
        styles['Normal'].textColor = colors.HexColor("#2D3748")
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 13
        
        # Custom styles
        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#1A365D"),
            spaceAfter=15
        )
        
        section_heading = ParagraphStyle(
            'SectionHeading',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#1A365D"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True
        )
        
        meta_label_style = ParagraphStyle(
            'MetaLabel',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#4A5568")
        )
        
        meta_value_style = ParagraphStyle(
            'MetaValue',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#1A202C")
        )
        
        table_cell_style = ParagraphStyle(
            'TableCell',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=9,
            leading=11
        )
        
        table_header_style = ParagraphStyle(
            'TableHeader',
            parent=styles['Normal'],
            fontName='Helvetica-Bold',
            fontSize=9,
            leading=11,
            textColor=colors.white
        )

        story = []
        
        # 1. Header Metadata Table (Deep Navy themed)
        meta_data = [
            [
                Paragraph("<b>Doc ID:</b>", meta_label_style), Paragraph(doc_data["doc_id"], meta_value_style),
                Paragraph("<b>Date:</b>", meta_label_style), Paragraph(doc_data["date"].strftime("%Y-%m-%d"), meta_value_style),
                Paragraph("<b>Revision:</b>", meta_label_style), Paragraph(doc_data["revision"], meta_value_style)
            ],
            [
                Paragraph("<b>Equipment Tag:</b>", meta_label_style), Paragraph(doc_data["equipment_tag"], meta_value_style),
                Paragraph("<b>Asset Name:</b>", meta_label_style), Paragraph(doc_data["equipment_name"], meta_value_style),
                Paragraph("<b>Risk Level:</b>", meta_label_style), Paragraph(f"<font color='red'><b>{doc_data['risk_level']}</b></font>" if doc_data['risk_level'] == 'Critical' else doc_data['risk_level'], meta_value_style)
            ]
        ]
        
        # 487 is printable width (595 - 108)
        meta_table = Table(meta_data, colWidths=[90, 80, 90, 80, 77, 70])
        meta_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E0")),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#EDF2F7")),
            ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#EDF2F7")),
            ('BACKGROUND', (4,0), (4,-1), colors.HexColor("#EDF2F7")),
            ('PADDING', (0,0), (-1,-1), 5),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        
        story.append(meta_table)
        story.append(Spacer(1, 15))
        
        # Title
        story.append(Paragraph(doc_data["title"], title_style))
        story.append(Spacer(1, 10))
        
        # 2. Observations
        story.append(Paragraph("1. Technical Observations & Description", section_heading))
        story.append(Paragraph(doc_data["observations"], styles['Normal']))
        story.append(Spacer(1, 10))
        
        # 3. Engineering Notes
        story.append(Paragraph("2. Engineering Notes & Work Instructions", section_heading))
        story.append(Paragraph(doc_data["engineering_notes"].replace('\n', '<br/>'), styles['Normal']))
        story.append(Spacer(1, 10))
        
        # 4. Structured Data Table
        if doc_data.get("tables"):
            story.append(Paragraph("3. Logged Parameters / Analytical Table", section_heading))
            raw_table_data = doc_data["tables"]
            table_data = []
            
            # Header Row
            table_data.append([Paragraph(f"<b>{col}</b>", table_header_style) for col in raw_table_data[0]])
            
            # Data Rows
            for row in raw_table_data[1:]:
                table_data.append([Paragraph(str(col), table_cell_style) for col in row])
                
            col_widths = [487 / len(raw_table_data[0])] * len(raw_table_data[0])
            data_table = Table(table_data, colWidths=col_widths)
            
            t_style = TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1A365D")),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E0")),
                ('PADDING', (0,0), (-1,-1), 6),
            ])
            # Add alternating row background
            for r_idx in range(1, len(table_data)):
                if r_idx % 2 == 0:
                    t_style.add('BACKGROUND', (0, r_idx), (-1, r_idx), colors.HexColor("#F7FAFC"))
                    
            data_table.setStyle(t_style)
            story.append(data_table)
            story.append(Spacer(1, 10))
            
        # 5. Recommendations
        story.append(Paragraph("4. Corrective Actions & Recommendations", section_heading))
        story.append(Paragraph(doc_data["recommendations"].replace('\n', '<br/>'), styles['Normal']))
        story.append(Spacer(1, 15))
        
        # 6. References (Knowledge Graph capability)
        if doc_data.get("references"):
            story.append(Paragraph("5. Related Documents & Traceability References", section_heading))
            ref_texts = []
            for ref in doc_data["references"]:
                ref_texts.append(f"• <b>Ref Doc:</b> {ref['doc_id']} — <b>Relationship:</b> {ref['relationship']}")
            story.append(Paragraph("<br/>".join(ref_texts), styles['Normal']))
            story.append(Spacer(1, 15))
            
        # 7. Approvals Matrix (Keep Together to avoid orphan signatures)
        app_data = [
            [
                Paragraph("<b>Prepared By:</b>", meta_label_style),
                Paragraph("<b>Reviewed By:</b>", meta_label_style),
                Paragraph("<b>Approved By:</b>", meta_label_style)
            ],
            [
                Paragraph(f"{doc_data['prepared_by']}<br/>Emp ID: {doc_data['prepared_by_id']}", meta_value_style),
                Paragraph(f"{doc_data['reviewed_by']}<br/>Emp ID: {doc_data['reviewed_by_id']}", meta_value_style),
                Paragraph(f"{doc_data['approved_by']}<br/>Emp ID: {doc_data['approved_by_id']}", meta_value_style)
            ],
            [
                Paragraph("Signature: ________________", meta_label_style),
                Paragraph("Signature: ________________", meta_label_style),
                Paragraph("Signature: ________________", meta_label_style)
            ]
        ]
        app_table = Table(app_data, colWidths=[162, 162, 163])
        app_table.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E0")),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#EDF2F7")),
            ('PADDING', (0,0), (-1,-1), 6),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ]))
        
        story.append(KeepTogether([
            Paragraph("6. Review & Approval Matrix", section_heading),
            app_table
        ]))
        
        # Build Document
        doc.build(story, canvasmaker=NumberedCanvas)
        
    def generate_docx(self, doc_data, category_name):
        category_docx_dir = self.docx_dir / category_name
        category_docx_dir.mkdir(parents=True, exist_ok=True)
        docx_path = category_docx_dir / f"{doc_data['doc_id']}.docx"
        
        doc = Document()
        
        # Setup Page Margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            
        # Title
        p_title = doc.add_paragraph()
        p_title_run = p_title.add_run(doc_data["title"])
        p_title_run.font.name = 'Arial'
        p_title_run.font.size = Pt(16)
        p_title_run.font.bold = True
        p_title_run.font.color.rgb = RGBColor(26, 54, 93)  # Deep Navy
        
        # Metadata Table
        table_meta = doc.add_table(rows=2, cols=6)
        table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        headers = [
            ("Doc ID:", doc_data["doc_id"]),
            ("Date:", doc_data["date"].strftime("%Y-%m-%d")),
            ("Revision:", doc_data["revision"]),
            ("Equipment Tag:", doc_data["equipment_tag"]),
            ("Asset Name:", doc_data["equipment_name"]),
            ("Risk Level:", doc_data["risk_level"])
        ]
        
        # Fill table metadata
        meta_mapping = [
            [(0,0), (0,1)], # Doc ID
            [(0,2), (0,3)], # Date
            [(0,4), (0,5)], # Revision
            [(1,0), (1,1)], # Tag
            [(1,2), (1,3)], # Name
            [(1,4), (1,5)]  # Risk
        ]
        
        for idx, mapping in enumerate(meta_mapping):
            lbl_pos, val_pos = mapping
            lbl, val = headers[idx]
            
            lbl_cell = table_meta.cell(lbl_pos[0], lbl_pos[1])
            lbl_cell.text = lbl
            set_cell_background(lbl_cell, "EDF2F7")
            set_cell_margins(lbl_cell, top=60, bottom=60, left=100, right=100)
            lbl_cell.paragraphs[0].runs[0].font.bold = True
            lbl_cell.paragraphs[0].runs[0].font.size = Pt(9)
            
            val_cell = table_meta.cell(val_pos[0], val_pos[1])
            val_cell.text = val
            set_cell_margins(val_cell, top=60, bottom=60, left=100, right=100)
            val_cell.paragraphs[0].runs[0].font.size = Pt(9)
            if lbl == "Risk Level:" and val == "Critical":
                val_cell.paragraphs[0].runs[0].font.bold = True
                val_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                
        doc.add_paragraph() # Spacing
        
        # Add Sections
        def add_section(num, title, content):
            p_sec = doc.add_paragraph()
            run_sec = p_sec.add_run(f"{num}. {title}")
            run_sec.font.name = 'Arial'
            run_sec.font.size = Pt(12)
            run_sec.font.bold = True
            run_sec.font.color.rgb = RGBColor(26, 54, 93)
            
            p_content = doc.add_paragraph()
            run_content = p_content.add_run(content)
            run_content.font.name = 'Calibri'
            run_content.font.size = Pt(10)
            
        add_section("1", "Technical Observations & Description", doc_data["observations"])
        add_section("2", "Engineering Notes & Work Instructions", doc_data["engineering_notes"])
        
        # Data Table
        if doc_data.get("tables"):
            p_sec = doc.add_paragraph()
            run_sec = p_sec.add_run("3. Logged Parameters / Analytical Table")
            run_sec.font.name = 'Arial'
            run_sec.font.size = Pt(12)
            run_sec.font.bold = True
            run_sec.font.color.rgb = RGBColor(26, 54, 93)
            
            raw_table_data = doc_data["tables"]
            table_data = doc.add_table(rows=len(raw_table_data), cols=len(raw_table_data[0]))
            table_data.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Header Row
            for c_idx, col_name in enumerate(raw_table_data[0]):
                cell = table_data.cell(0, c_idx)
                cell.text = col_name
                set_cell_background(cell, "1A365D")
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                cell.paragraphs[0].runs[0].font.size = Pt(9.5)
                
            # Data Rows
            for r_idx, row in enumerate(raw_table_data[1:]):
                for c_idx, val in enumerate(row):
                    cell = table_data.cell(r_idx + 1, c_idx)
                    cell.text = str(val)
                    set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                    cell.paragraphs[0].runs[0].font.size = Pt(9)
                    if (r_idx + 1) % 2 == 0:
                        set_cell_background(cell, "F7FAFC")
                        
            doc.add_paragraph() # Spacing
            
        # Recommendations
        add_section("4", "Corrective Actions & Recommendations", doc_data["recommendations"])
        
        # References
        if doc_data.get("references"):
            ref_texts = []
            for ref in doc_data["references"]:
                ref_texts.append(f"Ref Doc: {ref['doc_id']} — Relationship: {ref['relationship']}")
            add_section("5", "Related Documents & Traceability References", "\n".join(ref_texts))
            
        # Approvals Matrix
        p_sec = doc.add_paragraph()
        run_sec = p_sec.add_run("6. Review & Approval Matrix")
        run_sec.font.name = 'Arial'
        run_sec.font.size = Pt(12)
        run_sec.font.bold = True
        run_sec.font.color.rgb = RGBColor(26, 54, 93)
        
        table_app = doc.add_table(rows=3, cols=3)
        table_app.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        app_cols = [
            ("Prepared By:", f"{doc_data['prepared_by']}\nEmp ID: {doc_data['prepared_by_id']}"),
            ("Reviewed By:", f"{doc_data['reviewed_by']}\nEmp ID: {doc_data['reviewed_by_id']}"),
            ("Approved By:", f"{doc_data['approved_by']}\nEmp ID: {doc_data['approved_by_id']}")
        ]
        
        for c_idx, (lbl, val) in enumerate(app_cols):
            # Row 0: Labels
            cell_lbl = table_app.cell(0, c_idx)
            cell_lbl.text = lbl
            set_cell_background(cell_lbl, "EDF2F7")
            set_cell_margins(cell_lbl, top=60, bottom=60, left=100, right=100)
            cell_lbl.paragraphs[0].runs[0].font.bold = True
            cell_lbl.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Row 1: Names
            cell_val = table_app.cell(1, c_idx)
            cell_val.text = val
            set_cell_margins(cell_val, top=60, bottom=60, left=100, right=100)
            cell_val.paragraphs[0].runs[0].font.size = Pt(9)
            
            # Row 2: Signatures
            cell_sig = table_app.cell(2, c_idx)
            cell_sig.text = "Signature: ________________"
            set_cell_margins(cell_sig, top=60, bottom=60, left=100, right=100)
            cell_sig.paragraphs[0].runs[0].font.size = Pt(9.5)
            
        doc.save(str(docx_path))
